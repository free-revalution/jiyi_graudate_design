#!/usr/bin/env python3
"""
generate_thesis.py
Reads thesis_full.md and generates a formatted .docx thesis file
matching the exact formatting of the sample thesis (范文.docx).
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, "thesis_full.md")
OUTPUT_PATH = os.path.join(os.path.dirname(SCRIPT_DIR),
                           "基于Spring_Boot和Vue的智慧实验平台设计与实现.docx")

# ── Constants ────────────────────────────────────────────────────────────
# w:firstLine is in twips (1 twip = 1/20 pt = 1/1440 inch)
# 0.74 cm ≈ 420 twips  (0.74 / 2.54 * 1440)
FIRST_LINE_INDENT_TWIPS = 420
BODY_LINE_SPACING = 1.5

def _first_line_indent_xml():
    """Return an XML ind element for body first-line indent (2 chars)."""
    return parse_xml(
        f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
        f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
    )

# ── Helper: set run font ─────────────────────────────────────────────────
def set_run_font(run, cn_font, en_font, size_pt, bold=False, color=None):
    """Set Chinese + English font, size and bold on a run, also setting eastAsia."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)


def set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=None,
                          line_rule=None):
    """Set paragraph spacing via XML to avoid the limited python-docx API."""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    if line_spacing is not None:
        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
        rule = line_rule or WD_LINE_SPACING.MULTIPLE
        spacing.set(qn('w:lineRule'), 'auto' if rule == WD_LINE_SPACING.MULTIPLE else 'exact')


def add_page_break(doc):
    """Add a page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))
    # Make break paragraph invisible (zero height)
    set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=0)
    return para


def add_empty_line(doc, size_pt=12):
    """Add an empty line with given font size."""
    para = doc.add_paragraph()
    run = para.add_run()
    set_run_font(run, '宋体', 'Times New Roman', size_pt)
    return para


def set_section_page_layout(section):
    """Set A4 page layout with required margins."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.00)
    section.right_margin = Cm(3.00)
    section.gutter = Cm(0.20)


# ── Cover Page ───────────────────────────────────────────────────────────
def build_cover(doc):
    # "南 京 工 程 学 院"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('南 京 工 程 学 院')
    set_run_font(run, '楷体_GB2312', 'KaiTi_GB2312', 30)
    set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5)

    # Empty lines
    add_empty_line(doc, 30)
    add_empty_line(doc, 30)

    # "毕业设计说明书(论文)"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('毕业设计说明书(论文)')
    set_run_font(run, '黑体', 'SimHei', 45)
    set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5)

    # Empty lines
    for _ in range(4):
        add_empty_line(doc, 18)

    # Author info lines — 宋体 小三(15pt), CENTER
    info_lines = [
        ('作 者：', 'XXX'),
        ('学 号：', 'XXXXXXXXXX'),
        ('单 位：', '计算机工程学院'),
        ('专 业：', '软件工程'),
        ('题 目：', '基于Spring Boot和Vue的智慧实验平台设计与实现'),
    ]
    for label, value in info_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label + value)
        set_run_font(run, '宋体', 'Times New Roman', 15)
        set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5)
        pf = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="100" w:firstLine="200"/>'
        )
        para._element.get_or_add_pPr().append(pf)

    # Empty line
    add_empty_line(doc, 15)

    # 指导者 / 评阅者
    advisor_lines = [
        ('指导者：', 'XXX 教授'),
        ('评阅者：', 'XXX 副教授'),
    ]
    for label, value in advisor_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label + value)
        set_run_font(run, '宋体', 'Times New Roman', 15)
        set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5)
        pf = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="100" w:firstLine="200"/>'
        )
        para._element.get_or_add_pPr().append(pf)

    # Empty lines
    for _ in range(3):
        add_empty_line(doc, 18)

    # "2026年6月   南 京"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('2026年6月   南 京')
    set_run_font(run, '黑体', 'SimHei', 18)
    set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5)

    add_page_break(doc)


# ── Chinese Abstract ─────────────────────────────────────────────────────
def build_chinese_abstract(doc, lines):
    # Title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('毕业设计说明书（论文）中文摘要')
    set_run_font(run, '黑体', 'SimHei', 18, bold=True)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    # Body paragraphs
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('关键词：') or line.startswith('关键词:'):
            # Keywords line
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(line)
            set_run_font(run, '宋体', 'Times New Roman', 12)
            set_paragraph_spacing(para, before_pt=0, after_pt=0,
                                  line_spacing=BODY_LINE_SPACING)
        else:
            # Normal body paragraph
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(line)
            set_run_font(run, '宋体', 'Times New Roman', 12)
            set_paragraph_spacing(para, before_pt=0, after_pt=0,
                                  line_spacing=BODY_LINE_SPACING)
            para._element.get_or_add_pPr().append(_first_line_indent_xml())

    add_page_break(doc)


# ── English Abstract ─────────────────────────────────────────────────────
def build_english_abstract(doc, lines):
    # Title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('毕业设计说明书（论文）外文摘要')
    set_run_font(run, '黑体', 'SimHei', 18, bold=True)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        # First non-empty line is the English title
        if line.startswith('Key words:') or line.startswith('Keywords:'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(line)
            set_run_font(run, 'Times New Roman', 'Times New Roman', 12)
            set_paragraph_spacing(para, before_pt=0, after_pt=0,
                                  line_spacing=BODY_LINE_SPACING)
        elif not any(c > '一' for c in line) and len(line) > 20:
            # Check if this looks like a title (English, not too long)
            is_title = (lines.index(line.lstrip(), 0,
                        sum(1 for l in lines[:lines.index(line)] if l.strip())) == 0)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            set_run_font(run, 'Times New Roman', 'Times New Roman', 14, bold=True)
            set_paragraph_spacing(para, before_pt=6, after_pt=6, line_spacing=1.5)
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(line)
            set_run_font(run, 'Times New Roman', 'Times New Roman', 12)
            set_paragraph_spacing(para, before_pt=0, after_pt=0,
                                  line_spacing=BODY_LINE_SPACING)
            pf = parse_xml(
                f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
                f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
            )
            para._element.get_or_add_pPr().append(pf)

    add_page_break(doc)


# ── Table of Contents ────────────────────────────────────────────────────
def build_toc(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('目　　录')
    set_run_font(run, '黑体', 'SimHei', 16)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('[请在Word中插入自动目录]')
    set_run_font(run, '宋体', 'Times New Roman', 12)
    set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing=1.5)

    add_page_break(doc)


# ── Preface (前言) ───────────────────────────────────────────────────────
def build_preface(doc, lines):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('前  言')
    set_run_font(run, '黑体', 'SimHei', 16)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    for line in lines:
        line = line.strip()
        if not line or line == '---':
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(line)
        set_run_font(run, '宋体', 'Times New Roman', 12)
        set_paragraph_spacing(para, before_pt=0, after_pt=0,
                              line_spacing=BODY_LINE_SPACING)
        pf = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
            f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
        )
        para._element.get_or_add_pPr().append(pf)

    # No trailing page break here — build_chapter adds its own leading break


# ── Chapter ──────────────────────────────────────────────────────────────
def build_chapter(doc, chapter_title, lines):
    # Page break before each chapter
    add_page_break(doc)

    # Chapter title — 黑体 16pt CENTER
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(chapter_title)
    set_run_font(run, '黑体', 'SimHei', 16)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        # Horizontal rule
        if stripped == '---':
            continue

        # Code block toggle
        if stripped.startswith('```'):
            if in_code_block:
                # End code block — flush
                _flush_code(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            continue

        if not stripped:
            continue

        # Section title (## ...)
        if stripped.startswith('## ') and not stripped.startswith('### '):
            sec_title = stripped[3:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(sec_title)
            set_run_font(run, '黑体', 'SimHei', 14, bold=False)
            set_paragraph_spacing(para, before_pt=6, after_pt=6, line_spacing=1.5)
            continue

        # Subsection title (### ...)
        if stripped.startswith('### '):
            subsec_title = stripped[4:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(subsec_title)
            set_run_font(run, '宋体', 'SimHei', 14, bold=True)
            set_paragraph_spacing(para, before_pt=6, after_pt=6, line_spacing=1.5)
            continue

        # Figure / table placeholder [图...] or [表...]
        if stripped.startswith('[图') or stripped.startswith('[表'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            set_run_font(run, '宋体', 'Times New Roman', 12)
            set_paragraph_spacing(para, before_pt=6, after_pt=6, line_spacing=1.5)
            continue

        # Bold paragraph lead (e.g., **用户与权限相关表：**)
        # We handle inline bold markers
        if '**' in stripped:
            _add_rich_paragraph(doc, stripped)
            continue

        # Normal body paragraph
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(stripped)
        set_run_font(run, '宋体', 'Times New Roman', 12)
        set_paragraph_spacing(para, before_pt=0, after_pt=0,
                              line_spacing=BODY_LINE_SPACING)
        pf = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
            f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
        )
        para._element.get_or_add_pPr().append(pf)


def _flush_code(doc, code_lines):
    """Write collected code lines into a code block paragraph."""
    if not code_lines:
        return
    text = '\n'.join(code_lines)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, '宋体', 'Courier New', 10)
    set_paragraph_spacing(para, before_pt=3, after_pt=3, line_spacing=1.0)
    # Light gray shading
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'
    )
    para._element.get_or_add_pPr().append(shd)


def _add_rich_paragraph(doc, text):
    """Add a paragraph that may contain **bold** inline markers."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(para, before_pt=0, after_pt=0,
                          line_spacing=BODY_LINE_SPACING)
    pf = parse_xml(
        f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
        f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
    )
    para._element.get_or_add_pPr().append(pf)

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, '宋体', 'Times New Roman', 12, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, '宋体', 'Times New Roman', 12)


# ── References (参考文献) ───────────────────────────────────────────────
def build_references(doc, lines):
    add_page_break(doc)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('参考文献')
    set_run_font(run, '黑体', 'SimHei', 16)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped == '---':
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(stripped)
        set_run_font(run, '宋体', 'Times New Roman', 12)
        set_paragraph_spacing(para, before_pt=0, after_pt=0,
                              line_spacing=BODY_LINE_SPACING)


# ── Acknowledgement (致谢) ──────────────────────────────────────────────
def build_acknowledgement(doc, lines):
    add_page_break(doc)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('致  谢')
    set_run_font(run, '黑体', 'SimHei', 16)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped == '---':
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(stripped)
        set_run_font(run, '宋体', 'Times New Roman', 12)
        set_paragraph_spacing(para, before_pt=0, after_pt=0,
                              line_spacing=BODY_LINE_SPACING)
        pf = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
            f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
        )
        para._element.get_or_add_pPr().append(pf)


# ── English Abstract — improved ──────────────────────────────────────────
def build_english_abstract_v2(doc, title_line, body_lines, keywords_line):
    """Build English abstract with proper formatting."""
    # Section heading
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('毕业设计说明书（论文）外文摘要')
    set_run_font(run, '黑体', 'SimHei', 18, bold=True)
    set_paragraph_spacing(para, before_pt=12, after_pt=12, line_spacing=1.5)

    # English title — bold, centered
    if title_line:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title_line.strip())
        set_run_font(run, 'Times New Roman', 'Times New Roman', 14, bold=True)
        set_paragraph_spacing(para, before_pt=6, after_pt=6, line_spacing=1.5)

    # Body paragraphs
    for line in body_lines:
        stripped = line.strip()
        if not stripped:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(stripped)
        set_run_font(run, 'Times New Roman', 'Times New Roman', 12)
        set_paragraph_spacing(para, before_pt=0, after_pt=0,
                              line_spacing=BODY_LINE_SPACING)
        pf = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="200" '
            f'w:firstLine="{FIRST_LINE_INDENT_TWIPS}"/>'
        )
        para._element.get_or_add_pPr().append(pf)

    # Key words line
    if keywords_line:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(keywords_line.strip())
        set_run_font(run, 'Times New Roman', 'Times New Roman', 12)
        set_paragraph_spacing(para, before_pt=0, after_pt=0,
                              line_spacing=BODY_LINE_SPACING)

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════
# Main parsing & generation
# ══════════════════════════════════════════════════════════════════════════

def parse_markdown(md_text):
    """
    Parse the markdown into logical sections.
    Returns a list of (section_type, title, lines) tuples.
    """
    lines = md_text.split('\n')
    sections = []
    current_type = None
    current_title = None
    current_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Detect top-level heading (# ...)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading = stripped[2:].strip()

            # Save previous section
            if current_type is not None:
                sections.append((current_type, current_title, list(current_lines)))
                current_lines = []

            # Determine section type from heading
            if '中文摘要' in heading:
                current_type = 'chinese_abstract'
                current_title = heading
            elif '外文摘要' in heading:
                current_type = 'english_abstract'
                current_title = heading
            elif heading == '目录':
                current_type = 'toc'
                current_title = heading
            elif '前言' in heading:
                current_type = 'preface'
                current_title = heading
            elif heading.startswith('第') and '章' in heading:
                current_type = 'chapter'
                current_title = heading
            elif heading == '参考文献':
                current_type = 'references'
                current_title = heading
            elif '致谢' in heading:
                current_type = 'acknowledgement'
                current_title = heading
            else:
                current_type = 'chapter'
                current_title = heading
        else:
            current_lines.append(line)

        i += 1

    # Save last section
    if current_type is not None:
        sections.append((current_type, current_title, list(current_lines)))

    return sections


def generate():
    """Main generation function."""
    # Read markdown
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        md_text = f.read()

    sections = parse_markdown(md_text)

    # Create document
    doc = Document()

    # Set default page layout
    section = doc.sections[0]
    set_section_page_layout(section)

    # ── Cover ──
    build_cover(doc)

    # ── Process each section ──
    for sec_type, title, lines in sections:

        if sec_type == 'chinese_abstract':
            # Find keywords line vs body lines
            body_lines = []
            kw_line = None
            for ln in lines:
                s = ln.strip()
                if s.startswith('关键词：') or s.startswith('关键词:'):
                    kw_line = s
                elif s:
                    body_lines.append(s)
            build_chinese_abstract(doc, body_lines + ([kw_line] if kw_line else []))

        elif sec_type == 'english_abstract':
            # Parse: first non-empty line is title, last keywords line, rest is body
            non_empty = [ln.strip() for ln in lines if ln.strip()]
            title_line = None
            body_lines = []
            kw_line = None

            for idx, ln in enumerate(non_empty):
                if ln.startswith('Key words:') or ln.startswith('Keywords:'):
                    kw_line = ln
                elif idx == 0:
                    title_line = ln
                else:
                    body_lines.append(ln)

            build_english_abstract_v2(doc, title_line, body_lines, kw_line)

        elif sec_type == 'toc':
            build_toc(doc)

        elif sec_type == 'preface':
            build_preface(doc, lines)

        elif sec_type == 'chapter':
            build_chapter(doc, title, lines)

        elif sec_type == 'references':
            build_references(doc, lines)

        elif sec_type == 'acknowledgement':
            build_acknowledgement(doc, lines)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"Thesis generated: {OUTPUT_PATH}")


if __name__ == '__main__':
    generate()
